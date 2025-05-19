# ui/report_view.py
import sys
from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QPushButton, QLabel,
                             QFileDialog, QMessageBox, QLineEdit, QDateEdit,
                             QComboBox, QFormLayout, QHBoxLayout)
from PyQt5.QtSql import QSqlDatabase, QSqlQuery, QSqlError
from PyQt5.QtCore import Qt, QDate

# Импортируем библиотеку для работы с .docx
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Библиотека 'python-docx' не найдена. Установите ее: pip install python-docx")
    # Можно добавить обработку ошибки или уведомление пользователя в GUI

# Импортируем схему базы данных для получения названий таблиц и столбцов
from database import DATABASE_SCHEMA

class ReportView(QWidget):
    def __init__(self, db_connection):
        super().__init__()

        self.db = db_connection
        if not self.db or not self.db.isOpen():
            print("Ошибка: Соединение с базой данных не установлено или закрыто.")
            return

        self.setWindowTitle("Создание отчетов по инвентаризации")

        self.layout = QVBoxLayout(self)

        info_label = QLabel("Выберите параметры для отчета по инвентаризации и нажмите 'Сформировать отчет'.")
        self.layout.addWidget(info_label)

        # --- Элементы управления для фильтрации (адаптированы к новой БД) ---
        filter_layout = QFormLayout()

        # Фильтр по Категории
        self.category_combo = QComboBox()
        self.category_combo.addItem("Все категории", None) # Опция "Все категории"
        self._populate_category_combo() # Заполняем комбобокс категориями

        # Фильтр по Подкатегории (будет зависеть от выбранной категории)
        self.subcategory_combo = QComboBox()
        self.subcategory_combo.addItem("Все подкатегории", None) # Опция "Все подкатегории"
        # TODO: Реализовать обновление этого комбобокса при изменении category_combo
        # self.category_combo.currentIndexChanged.connect(self._populate_subcategory_combo)
        self._populate_subcategory_combo() # Заполняем изначально все подкатегории

        # Фильтр по Типу единицы
        self.unit_type_combo = QComboBox()
        self.unit_type_combo.addItem("Все типы единиц", None)
        self._populate_unit_type_combo()

        # Фильтр по Статусу заказа
        self.order_status_combo = QComboBox()
        self.order_status_combo.addItem("Все статусы", None)
        self._populate_order_status_combo()


        self.cabinet_input = QLineEdit()
        self.cabinet_input.setPlaceholderText("Фильтр по кабинету (опционально)")

        self.manufacturer_input = QLineEdit()
        self.manufacturer_input.setPlaceholderText("Фильтр по производителю (опционально)")

        self.model_input = QLineEdit()
        self.model_input.setPlaceholderText("Фильтр по модели (опционально)")

        self.serial_number_input = QLineEdit()
        self.serial_number_input.setPlaceholderText("Фильтр по серийному номеру (опционально)")

        self.inventory_number_input = QLineEdit()
        self.inventory_number_input.setPlaceholderText("Фильтр по инвентарному номеру (опционально)")


        date_range_layout = QHBoxLayout()
        self.start_date_edit = QDateEdit(calendarPopup=True)
        self.start_date_edit.setDate(QDate(2000, 1, 1)) # Пример начальной даты
        self.start_date_edit.setDisplayFormat("yyyy-MM-dd")

        date_range_layout.addWidget(QLabel("Дата заказа (с):"))
        date_range_layout.addWidget(self.start_date_edit)

        self.end_date_edit = QDateEdit(calendarPopup=True)
        self.end_date_edit.setDate(QDate.currentDate()) # Текущая дата по умолчанию
        self.end_date_edit.setDisplayFormat("yyyy-MM-dd")

        date_range_layout.addWidget(QLabel("по:"))
        date_range_layout.addWidget(self.end_date_edit)
        date_range_layout.addStretch() # Растягиваем

        filter_layout.addRow("Категория:", self.category_combo)
        filter_layout.addRow("Подкатегория:", self.subcategory_combo)
        filter_layout.addRow("Тип единицы:", self.unit_type_combo)
        filter_layout.addRow("Статус заказа:", self.order_status_combo)
        filter_layout.addRow("Кабинет:", self.cabinet_input)
        filter_layout.addRow("Производитель:", self.manufacturer_input)
        filter_layout.addRow("Модель:", self.model_input)
        filter_layout.addRow("Серийный номер:", self.serial_number_input)
        filter_layout.addRow("Инвентарный номер:", self.inventory_number_input)
        filter_layout.addRow(date_range_layout) # Добавляем макет с датами

        self.layout.addLayout(filter_layout)
        self.layout.addStretch() # Растягиваем, чтобы кнопка была внизу

        # --- Кнопка формирования отчета ---
        generate_button = QPushButton("Сформировать отчет (.docx)")
        generate_button.clicked.connect(self._generate_report)
        self.layout.addWidget(generate_button)

    # --- Методы заполнения комбобоксов ---

    def _populate_category_combo(self):
        """Заполняет QComboBox категориями из базы данных."""
        self.category_combo.clear()
        self.category_combo.addItem("Все категории", None) # Опция "Все"
        query = QSqlQuery("SELECT id_category, category FROM Category ORDER BY category", self.db)
        while query.next():
            item_id = query.value(0)
            item_name = query.value(1)
            self.category_combo.addItem(item_name, item_id) # Сохраняем ID как UserData

    def _populate_subcategory_combo(self, category_id=None):
        """Заполняет QComboBox подкатегориями, опционально фильтруя по категории."""
        self.subcategory_combo.clear()
        self.subcategory_combo.addItem("Все подкатегории", None) # Опция "Все"
        query = QSqlQuery(self.db)
        sql = "SELECT id_subcategory, subcategory FROM Subcategory"
        params = []
        if category_id is not None:
             sql += " WHERE id_category = ?"
             params.append(category_id)
        sql += " ORDER BY subcategory"

        query.prepare(sql)
        for param in params:
             query.addBindValue(param)

        if query.exec_():
            while query.next():
                item_id = query.value(0)
                item_name = query.value(1)
                self.subcategory_combo.addItem(item_name, item_id)
        else:
            print("Ошибка при загрузке подкатегорий:", query.lastError().text())


    def _populate_unit_type_combo(self):
        """Заполняет QComboBox типами единиц из базы данных."""
        self.unit_type_combo.clear()
        self.unit_type_combo.addItem("Все типы единиц", None) # Опция "Все"
        query = QSqlQuery("SELECT id_unit_type, unit_type FROM Unit_type ORDER BY unit_type", self.db)
        while query.next():
            item_id = query.value(0)
            item_name = query.value(1)
            self.unit_type_combo.addItem(item_name, item_id) # Сохраняем ID как UserData

    def _populate_order_status_combo(self):
        """Заполняет QComboBox статусами заказов из базы данных."""
        self.order_status_combo.clear()
        self.order_status_combo.addItem("Все статусы", None) # Опция "Все"
        query = QSqlQuery("SELECT id_order_status, order_status FROM Order_status ORDER BY order_status", self.db)
        while query.next():
            item_id = query.value(0)
            item_name = query.value(1)
            self.order_status_combo.addItem(item_name, item_id) # Сохраняем ID как UserData


    def _generate_report(self):
        """Формирует отчет на основе данных из базы и сохраняет его в .docx."""
        print("Формирование отчета...")

        # Получаем выбранные параметры фильтрации
        selected_category_id = self.category_combo.currentData()
        selected_subcategory_id = self.subcategory_combo.currentData()
        selected_unit_type_id = self.unit_type_combo.currentData()
        selected_order_status_id = self.order_status_combo.currentData()
        cabinet_filter = self.cabinet_input.text().strip()
        manufacturer_filter = self.manufacturer_input.text().strip()
        model_filter = self.model_input.text().strip()
        serial_number_filter = self.serial_number_input.text().strip()
        inventory_number_filter = self.inventory_number_input.text().strip()
        start_date = self.start_date_edit.date().toString(Qt.ISODate)
        end_date = self.end_date_edit.date().toString(Qt.ISODate)

        # Формируем SQL-запрос с учетом фильтров (используем новые таблицы)
        query_string = """
            SELECT
                ui.inventory_number,
                ui.serial_number,
                ui.manufacturer,
                ui.model,
                c.category,       -- Название категории из связанной таблицы
                sc.subcategory,   -- Название подкатегории из связанной таблицы
                ut.unit_type,     -- Название типа единицы из связанной таблицы
                ui.cabinet,
                os.order_status,  -- Название статуса заказа из связанной таблицы
                ui.date_order_buhgaltery,
                ui.date_issue,
                ui.notice,
                uei.device_name,  -- Данные из расширенной информации (если есть)
                uei.ip,
                uei.mac
            FROM
                Units_inventory ui
            LEFT JOIN Category c ON ui.id_category = c.id_category
            LEFT JOIN Subcategory sc ON ui.id_subcategory = sc.id_subcategory
            LEFT JOIN Unit_type ut ON ui.id_unit_type = ut.id_unit_type
            LEFT JOIN Order_status os ON ui.id_order_status = os.id_order_status
            LEFT JOIN Units_extended_info uei ON ui.id_unit_inventory = uei.id_unit_inventory -- Связь 1-к-1
            WHERE 1=1
        """
        query_params = []

        if selected_category_id is not None:
            query_string += " AND ui.id_category = ?"
            query_params.append(selected_category_id)

        if selected_subcategory_id is not None:
            query_string += " AND ui.id_subcategory = ?"
            query_params.append(selected_subcategory_id)

        if selected_unit_type_id is not None:
            query_string += " AND ui.id_unit_type = ?"
            query_params.append(selected_unit_type_id)

        if selected_order_status_id is not None:
            query_string += " AND ui.id_order_status = ?"
            query_params.append(selected_order_status_id)

        if cabinet_filter:
            query_string += " AND ui.cabinet LIKE ?"
            query_params.append(f"%{cabinet_filter}%")

        if manufacturer_filter:
            query_string += " AND ui.manufacturer LIKE ?"
            query_params.append(f"%{manufacturer_filter}%")

        if model_filter:
            query_string += " AND ui.model LIKE ?"
            query_params.append(f"%{model_filter}%")

        if serial_number_filter:
            query_string += " AND ui.serial_number LIKE ?"
            query_params.append(f"%{serial_number_filter}%")

        if inventory_number_filter:
            query_string += " AND ui.inventory_number LIKE ?"
            query_params.append(f"%{inventory_number_filter}%")


        # Фильтр по дате заказа (date_order_buhgaltery)
        query_string += " AND ui.date_order_buhgaltery BETWEEN ? AND ?"
        query_params.append(start_date)
        query_params.append(end_date)

        query_string += " ORDER BY ui.date_order_buhgaltery, ui.inventory_number" # Сортировка

        query = QSqlQuery(self.db)
        query.prepare(query_string)
        for param in query_params:
            query.addBindValue(param)

        if not query.exec_():
            error_message = f"Ошибка при выполнении запроса к базе данных:\n{query.lastError().text()}"
            print(error_message)
            QMessageBox.critical(self, "Ошибка базы данных", error_message)
            return

        # Создаем новый документ Word
        try:
            document = Document()

            # Добавляем заголовок
            document.add_heading('Отчет по инвентаризации', 0)

            # Добавляем информацию о фильтрах (опционально)
            document.add_paragraph(f"Сформирован: {QDate.currentDate().toString(Qt.ISODate)}")
            # TODO: Добавить более подробную информацию о примененных фильтрах
            document.add_paragraph("Примененные фильтры...")
            document.add_paragraph("") # Пустая строка для отступа

            # Добавляем таблицу
            # Определяем заголовки столбцов для отчета (соответствуют SELECT в запросе)
            headers = [
                "Инв. номер", "Сер. номер", "Производитель", "Модель",
                "Категория", "Подкатегория", "Тип единицы", "Кабинет",
                "Статус заказа", "Дата заказа", "Дата выдачи", "Примечание",
                "Имя устройства", "IP", "MAC"
            ]
            table = document.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid' # Применяем стиль сетки

            # Заполняем заголовки таблицы
            header_cells = table.rows[0].cells
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text

            # Заполняем таблицу данными из запроса
            row_count = 0
            while query.next():
                row_cells = table.add_row().cells
                # Заполняем ячейки данными из запроса по индексу
                for i in range(len(headers)):
                    value = query.value(i)
                    # Преобразуем QDate в строку, если это дата
                    if isinstance(value, QDate):
                         row_cells[i].text = value.toString(Qt.ISODate)
                    elif value is None:
                         row_cells[i].text = "" # Пустая строка для NULL
                    else:
                         row_cells[i].text = str(value)

                row_count += 1

            if row_count == 0:
                 document.add_paragraph("Нет данных, соответствующих выбранным фильтрам.")


            # Диалог сохранения файла
            default_filename = f"Отчет_инвентаризация_{QDate.currentDate().toString('yyyyMMdd')}.docx"
            file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить отчет",
                                                       default_filename,
                                                       "Документы Word (*.docx);;Все файлы (*)")

            if file_path:
                document.save(file_path)
                QMessageBox.information(self, "Отчет сформирован", f"Отчет успешно сохранен в:\n{file_path}")
                print(f"Отчет сохранен: {file_path}")
            else:
                print("Сохранение отчета отменено.")

        except ImportError:
             QMessageBox.critical(self, "Ошибка", "Библиотека 'python-docx' не установлена. Пожалуйста, установите ее (`pip install python-docx`).")
        except Exception as e:
            error_message = f"Произошла ошибка при создании или сохранении отчета:\n{e}"
            print(error_message)
            QMessageBox.critical(self, "Ошибка", error_message)

